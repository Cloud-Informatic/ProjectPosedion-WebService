from cryptography.hazmat.backends import default_backend
from cryptography.hazmat.primitives.asymmetric import rsa
from hybrid_rsa_aes import HybridCipher
from docx import Document

#Public ve private key
rsa_private_key = rsa.generate_private_key(
    public_exponent=65537, key_size=2048, backend=default_backend()
)
rsa_public_key = rsa_private_key.public_key()

#Docx Baslangic
d = open('hibrit_deneme.docx', 'rb')
document = Document(d)
d.close()

encrypt_docx = HybridCipher().encrypt(rsa_public_key=rsa_public_key, data=str(document))
document1 = Document()
paragraph = document1.add_paragraph(encrypt_docx)
document1.save("enc_hibrit_deneme.docx")
"""
with open('enc_hibrit_deneme.docx', 'rb') as encrypted_file:
    encrypted = encrypted_file.read()

decrypt_message = HybridCipher().decrypt(
    rsa_private_key=rsa_private_key, cipher_text=encrypted
)

document2 = Document()
paragraph = document2.add_paragraph(decrypt_message)
document2.save("dec_hibrit_deneme.docx")
"""

document2 = Document()
decrypt_docx = HybridCipher().decrypt(
    rsa_private_key=rsa_private_key, cipher_text=encrypt_docx
)
paragraph = document2.add_paragraph(decrypt_docx)
document2.save("dec_hibrit_deneme.docx")
#Docx Bitis

#Txt Baslangic
dosya="hibrit_deneme.txt"
with open(dosya,"rb") as file:
    veri=file.read().splitlines()
encrypt_txt = HybridCipher().encrypt(rsa_public_key=rsa_public_key, data=str(veri))

f=open("enc_hibrit_deneme.txt","w+")
with open("enc_hibrit_deneme.txt","w+") as file:
    file.write(encrypt_txt)

decrypt_txt = HybridCipher().decrypt(
    rsa_private_key=rsa_private_key, cipher_text=encrypt_txt
)

f=open("dec_hibrit_deneme.txt","w+")
f.write(decrypt_txt)
f.close()
#Txt Bitis
#Devam edilecek...
